#!/usr/bin/env python3
"""Run structural and image-integrity checks on a review-manuscript DOCX."""

from __future__ import annotations

import argparse
import hashlib
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document


FIGURE_RE = re.compile(r"^\s*Fig\.\s*(\d+)\b", re.IGNORECASE)
TABLE_RE = re.compile(r"^\s*Table\s+(\d+)\b", re.IGNORECASE)
REFERENCE_RE = re.compile(r"^\s*\[(\d+)\]\s+")
CITATION_RE = re.compile(r"\[(\d+(?:\s*[-,]\s*\d+)*)\]")
FORBIDDEN = (
    "how to use this draft",
    "working manuscript",
    "author-only",
    "version 1.0",
    "minimum auditable checkpoints",
)


def is_sequential(values: list[int]) -> bool:
    return values == list(range(1, len(values) + 1))


def expand_citation_group(group: str) -> set[int]:
    result: set[int] = set()
    for part in re.split(r"\s*,\s*", group):
        if "-" in part:
            start_text, end_text = re.split(r"\s*-\s*", part, maxsplit=1)
            start, end = int(start_text), int(end_text)
            if start <= end and end - start <= 500:
                result.update(range(start, end + 1))
        else:
            result.add(int(part))
    return result


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Audit DOCX figures, tables, references, citations, and embedded images."
    )
    parser.add_argument("docx", type=Path, help="DOCX manuscript")
    args = parser.parse_args()

    document = Document(args.docx)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    full_text = "\n".join(paragraphs)

    figures = [
        int(match.group(1))
        for text in paragraphs
        if (match := FIGURE_RE.match(text))
    ]
    tables = [
        int(match.group(1))
        for text in paragraphs
        if (match := TABLE_RE.match(text))
    ]
    references = [
        int(match.group(1))
        for text in paragraphs
        if (match := REFERENCE_RE.match(text))
    ]

    references_start = next(
        (index for index, text in enumerate(paragraphs) if text.lower() == "references"),
        len(paragraphs),
    )
    body_text = "\n".join(paragraphs[:references_start])
    cited: set[int] = set()
    for group in CITATION_RE.findall(body_text):
        cited.update(expand_citation_group(group))

    duplicate_media: dict[str, list[str]] = {}
    with zipfile.ZipFile(args.docx) as archive:
        media_names = sorted(
            name for name in archive.namelist() if name.startswith("word/media/")
        )
        by_hash: dict[str, list[str]] = {}
        for name in media_names:
            digest = hashlib.sha256(archive.read(name)).hexdigest()
            by_hash.setdefault(digest, []).append(name)
        duplicate_media = {
            digest: names for digest, names in by_hash.items() if len(names) > 1
        }

    errors: list[str] = []
    warnings: list[str] = []

    if figures and not is_sequential(figures):
        errors.append(f"figure captions are not sequential: {figures}")
    if tables and not is_sequential(tables):
        errors.append(f"table captions are not sequential: {tables}")
    if references and not is_sequential(references):
        errors.append("reference list is not sequential")

    reference_set = set(references)
    missing_references = sorted(cited - reference_set)
    uncited_references = sorted(reference_set - cited)
    if missing_references:
        errors.append(f"in-text citations without reference entries: {missing_references}")
    if uncited_references:
        warnings.append(f"reference entries not detected in body citations: {uncited_references}")

    lower_text = full_text.casefold()
    for phrase in FORBIDDEN:
        if phrase in lower_text:
            warnings.append(f"possible internal drafting text: '{phrase}'")

    if duplicate_media:
        for names in duplicate_media.values():
            errors.append("duplicate embedded image bytes: " + ", ".join(names))

    caption_counts = Counter(figures)
    repeated_captions = sorted(number for number, count in caption_counts.items() if count > 1)
    if repeated_captions:
        errors.append(f"repeated figure caption numbers: {repeated_captions}")

    print(
        f"paragraphs={len(paragraphs)} tables={len(document.tables)} "
        f"inline_shapes={len(document.inline_shapes)}"
    )
    print(
        f"figure_captions={len(figures)} table_captions={len(tables)} "
        f"references={len(references)} cited_numbers={len(cited)} "
        f"embedded_media={len(media_names)}"
    )
    print(f"errors={len(errors)} warnings={len(warnings)}")
    for item in errors:
        print("ERROR", item)
    for item in warnings:
        print("WARNING", item)

    if duplicate_media:
        print("NOTE Exact-byte duplicate detection does not replace visual near-duplicate review.")
    return 1 if errors else 0


if __name__ == "__main__":
    sys.exit(main())

